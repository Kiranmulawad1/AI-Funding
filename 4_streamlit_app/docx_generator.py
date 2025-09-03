# # docx_generator.py

# import io
# from docx import Document

# def build_draft_prompt(profile, metadata):
#     return f"""
# You are a professional grant writer AI assistant.

# Use the following company profile and the funding program details to generate a complete funding application draft.

# ## Company Profile
# - Company Name: {profile.get("company_name", "Not specified")}
# - Location: {profile.get("location", "Not specified")}
# - Industry: {profile.get("industry", "Not specified")}
# - Goals: {profile.get("goals", "Not specified")}
# - Project Idea: {profile.get("project_idea", "Not specified")}
# - Funding Need: {profile.get("funding_need", "Not specified")}

# ## Funding Program
# - Name: {metadata.get("name", "Not specified")}
# - Amount: {metadata.get("amount", "Not specified")}
# - Deadline: {metadata.get("deadline", "Not specified")}
# - Eligibility: {metadata.get("eligibility", "Not specified")}

# ## Format Output as:
# 1. Title
# 2. Executive Summary
# 3. Objectives
# 4. Innovation
# 5. Budget Estimate
# 6. Relevance to Program
# 7. Contact Details

# Be professional and concise, but clear. Format the output as a ready-to-edit funding application.
# """

# def generate_funding_draft(metadata, profile, llm_client) -> io.BytesIO:
#     prompt = build_draft_prompt(profile, metadata)
    
#     response = llm_client.chat.completions.create(
#         model="gpt-4",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     draft_text = response.choices[0].message.content.strip()

#     doc = Document()
#     doc.add_heading("Funding Application Draft", 0)

#     for para in draft_text.split("\n\n"):
#         if para.strip():
#             doc.add_paragraph(para.strip())

#     # Save to in-memory buffer
#     bio = io.BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

# docx_generator.py
import io
from docx import Document

def build_draft_prompt(profile, metadata):
    """Build prompt for GPT to generate funding application draft"""
    return f"""You are a professional grant writer AI assistant.

Use the following company profile and the funding program details to generate a complete funding application draft.

## Company Profile
- Company Name: {profile.get("company_name", "Not specified")}
- Location: {profile.get("location", "Not specified")}
- Industry: {profile.get("industry", "Not specified")}
- Goals: {profile.get("goals", "Not specified")}
- Project Idea: {profile.get("project_idea", "Not specified")}
- Funding Need: {profile.get("funding_need", "Not specified")}

## Additional Details (if available)
- Technical Approach: {profile.get("technical_approach", "Not specified")}
- Market Opportunity: {profile.get("market_opportunity", "Not specified")}
- Timeline: {profile.get("timeline", "Not specified")}
- Expected Outcomes: {profile.get("expected_outcomes", "Not specified")}

## Funding Program
- Name: {metadata.get("name", "Not specified")}
- Domain: {metadata.get("domain", "Not specified")}
- Amount: {metadata.get("amount", "Not specified")}
- Deadline: {metadata.get("deadline", "Not specified")}
- Eligibility: {metadata.get("eligibility", "Not specified")}
- Location: {metadata.get("location", "Not specified")}

## Instructions
Create a professional funding application with these sections:

1. **Executive Summary** (2-3 paragraphs)
2. **Project Description** (detailed explanation of the project)
3. **Technical Approach** (methodology and innovation)
4. **Market Opportunity** (problem being solved, target market)
5. **Implementation Timeline** (key milestones and phases)
6. **Budget Overview** (high-level budget breakdown)
7. **Expected Outcomes** (deliverables and impact)
8. **Team Qualifications** (brief team overview)
9. **Relevance to Funding Program** (how it fits the program goals)

Be professional, specific, and compelling. Focus on innovation, feasibility, and impact."""

def generate_funding_draft(metadata, profile, llm_client) -> bytes:
    """Generate a funding application draft as a Word document"""
    
    try:
        # Build the prompt
        prompt = build_draft_prompt(profile, metadata)
        
        # Get GPT response
        response = llm_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2500
        )
        
        draft_text = response.choices[0].message.content.strip()
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = f"Funding Application: {metadata.get('name', 'Unknown Program')}"
        doc.add_heading(title, 0)
        
        # Add company info
        doc.add_heading("Company Information", level=1)
        company_info = f"""
Company: {profile.get('company_name', 'Not specified')}
Location: {profile.get('location', 'Not specified')}
Industry: {profile.get('industry', 'Not specified')}
        """.strip()
        doc.add_paragraph(company_info)
        
        # Add funding program info
        doc.add_heading("Funding Program Details", level=1)
        program_info = f"""
Program: {metadata.get('name', 'Not specified')}
Amount: {metadata.get('amount', 'Not specified')}
Deadline: {metadata.get('deadline', 'Not specified')}
        """.strip()
        doc.add_paragraph(program_info)
        
        # Add the generated application content
        doc.add_heading("Application Content", level=1)
        
        # Split the draft text into paragraphs and add them
        paragraphs = draft_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if it's a heading (starts with #, **, or a number followed by .)
                if (para.strip().startswith('#') or 
                    para.strip().startswith('**') or 
                    (len(para.split('.')) > 1 and para.split('.')[0].strip().isdigit())):
                    # Remove markdown formatting
                    heading_text = para.strip().replace('#', '').replace('**', '').strip()
                    doc.add_heading(heading_text, level=2)
                else:
                    doc.add_paragraph(para.strip())
        
        # Add footer
        doc.add_page_break()
        doc.add_heading("Contact Information", level=1)
        doc.add_paragraph("This application was generated using AI Grant Finder.")
        doc.add_paragraph(f"Generated on: {metadata.get('generated_date', 'Unknown date')}")
        
        # Save to in-memory buffer
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return bio.getvalue()
        
    except Exception as e:
        print(f"Error generating funding draft: {e}")
        
        # Create a simple fallback document
        doc = Document()
        doc.add_heading("Funding Application Draft", 0)
        doc.add_paragraph("An error occurred while generating the detailed application.")
        doc.add_paragraph(f"Error: {str(e)}")
        doc.add_paragraph(f"Company: {profile.get('company_name', 'Unknown')}")
        doc.add_paragraph(f"Program: {metadata.get('name', 'Unknown')}")
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return bio.getvalue()