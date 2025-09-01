import os
import re
import io
import html
import json
import pandas as pd
import numpy as np
import streamlit as st
from dateutil import parser
from dotenv import load_dotenv
from openai import OpenAI
from pinecone import Pinecone
import fitz  # PyMuPDF
from urllib.parse import unquote

# ---- DOCX import (robust) ----
try:
    from docx import Document  # pip install python-docx
    HAS_DOCX = True
    DOCX_IMPORT_ERR = None
except Exception as _e:
    HAS_DOCX = False
    DOCX_IMPORT_ERR = _e

# =========================
# 0) Setup & Environment
# =========================
load_dotenv()

OPENAI_API_KEY   = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENV     = os.getenv("PINECONE_ENV")
INDEX_NAME       = os.getenv("PINECONE_INDEX_NAME", "funding-search")
NAMESPACE        = os.getenv("PINECONE_NAMESPACE", "openai-v3")
FUNDING_CSV_PATH = os.getenv("FUNDING_CSV_PATH", "./data/merged_funding_data.csv")

required = {
    "OPENAI_API_KEY": OPENAI_API_KEY,
    "PINECONE_API_KEY": PINECONE_API_KEY,
    "PINECONE_ENV": PINECONE_ENV,
}
st.set_page_config(page_title="Smart Funding Chatbot", layout="centered")
if missing := [k for k, v in required.items() if not v]:
    st.error(f"Missing environment variables: {', '.join(missing)}")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)
pc = Pinecone(api_key=PINECONE_API_KEY, environment=PINECONE_ENV)

@st.cache_resource(show_spinner=False)
def get_index(_pc, name: str):
    return _pc.Index(name)

@st.cache_data(show_spinner=False)
def load_full_df(path: str):
    if os.path.exists(path):
        return pd.read_csv(path)
    return pd.DataFrame()

index = get_index(pc, INDEX_NAME)
df_full = load_full_df(FUNDING_CSV_PATH)

# =========================
# 1) Session State
# =========================
def _ensure_state():
    defaults = {
        "chat_history": [],
        "draft_blobs": {},
        "generated_query": "",
        "send_now": False,
        "last_query": None,
        "last_matches": None,     # list[dict]
        "last_chosen_ids": None,  # list[int] (1-based ids within last_matches)
        "last_reasons": None,     # {id: why}
        "last_enrich": None,      # {id: {"brief":..., "next_steps":[...]}}
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def _reset_session():
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    _ensure_state()
    st.rerun()

_ensure_state()

# =========================
# 2) Sidebar: PDF upload + Reset
# =========================
st.sidebar.subheader("📄 Upload Company Profile (Optional)")
uploaded_pdf = st.sidebar.file_uploader("Upload PDF", type=["pdf"])
if uploaded_pdf:
    def extract_text_from_pdf(uploaded_file):
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join(page.get_text() for page in doc).strip()

    def summarize_text_with_gpt(raw_text: str) -> str:
        prompt = f"""
Summarize this company profile into 2–3 lines for public funding discovery in Germany.
Focus on domain, goals, and funding need.

---
{raw_text[:6000]}
---
"""
        res = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return res.choices[0].message.content.strip()

    raw_text = extract_text_from_pdf(uploaded_pdf)
    st.session_state.generated_query = summarize_text_with_gpt(raw_text)
    st.sidebar.success("✅ Profile processed. Using it to search.")

st.sidebar.button("🔄 Reset conversation", on_click=_reset_session)

# =========================
# 3) Retrieval / Scoring helpers
# =========================
def get_embedding(text: str):
    return client.embeddings.create(input=[text], model="text-embedding-3-small").data[0].embedding

def safe_parse_deadline(deadline_str):
    try:
        if pd.isna(deadline_str) or str(deadline_str).strip() == "":
            return None
        return parser.parse(str(deadline_str), dayfirst=True, fuzzy=True)
    except Exception:
        return None

def fmt(x, fallback="Not specified"):
    if x is None:
        return fallback
    try:
        import math
        if isinstance(x, float) and math.isnan(x):
            return fallback
    except Exception:
        pass
    s = str(x).strip()
    if not s:
        return fallback
    lower = s.lower()
    if any(phrase in lower for phrase in (
        "information not found", "not specified", "n/a", "na", "none", "null"
    )):
        return fallback
    return s

def two_sentences(text, max_sentences=2):
    s = str(text or "").strip()
    if not s or s.lower() == "not specified":
        return "Not specified"
    import re
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'[.…]+$', '', s).strip()
    parts = re.split(r'(?<=[.!?])\s+', s)
    parts = [p.strip() for p in parts if p.strip()]
    out = ' '.join(parts[:max_sentences]) if parts else "Not specified"
    if out and out[-1] not in ".!?":
        out += "."
    return out

def compute_relevance_score(row, query, funding_need=0, target_domain="", user_location=""):
    get = (row.get if isinstance(row, dict) else lambda k, d=None: row.get(k, d))
    score = 0.0
    if target_domain and target_domain.lower() in str(get("domain","")).lower():
        score += 0.4
    try:
        nums = re.findall(r"\d+", str(get("amount","0")))
        amount_val = int(nums[-1]) if nums else 0
        if funding_need and amount_val >= funding_need:
            score += 0.3
    except Exception:
        pass
    dd = get("deadline_date")
    if isinstance(dd, pd.Timestamp) and pd.notnull(dd) and dd >= pd.Timestamp.now(tz="UTC"):
        score += 0.2
    tokens = re.findall(r"\w+", query.lower())
    if any(tok in str(get("description","")).lower() for tok in tokens):
        score += 0.1
    if user_location and user_location.lower() in str(get("location","")).lower():
        score += 0.1
    return round(score * 100)

def deadline_with_badge(m):
    dl = fmt(m.get("deadline"))
    dl_days = m.get("days_left")
    if dl != "Not specified":
        try:
            if isinstance(dl_days, (int, float)) and not pd.isna(dl_days):
                return f"{dl} (🕒 {int(dl_days)} days left)"
        except Exception:
            pass
        return dl
    return "Not specified"

def steps_from_dataset(m):
    steps = []
    url = (m.get("url") or "").strip()
    elig = fmt(m.get("eligibility"))
    proc = fmt(m.get("procedure"))
    deadline = fmt(m.get("deadline"))

    # Always inject the "official page" step when URL exists
    if url and url != "Not specified":
        steps.append(f'Visit the <a href="{html.escape(url)}" target="_blank">official page</a>')
    if elig and elig != "Not specified":
        steps.append("Confirm you meet eligibility requirements")
    if proc and proc != "Not specified":
        steps.append("Follow the described application procedure")
    if deadline and deadline != "Not specified" and len(steps) < 3:
        steps.append(f"Note the deadline: {deadline}")

    return steps[:3] if steps else [
        "Visit the official page",
        "Prepare a 1–2 page project summary & budget",
        "Contact the program office for clarifications",
    ]

# --- program name normalization ---
ACRONYMS = ("AI","R&D","EU","BMBF","EFRE","ERDF","SME","ML","KI")

def _fix_acronyms(s: str) -> str:
    for ac in ACRONYMS:
        s = re.sub(rf"\b{ac.title()}\b", ac, s)
    return s

def normalize_program_title(candidate: str, url: str = "") -> str:
    s = str(candidate or "").strip()
    if not s:
        return ""
    # looks like slug or has extension?
    if s.lower().endswith((".html", ".htm", ".php")) or re.fullmatch(r"[a-z0-9\-\._]+", s.lower()):
        s = re.sub(r"\.(html?|php)$", "", s, flags=re.I)
        s = s.replace("-", " ").replace("_", " ")
        s = unquote(s)
        s = re.sub(r"\s+", " ", s).strip().title()
        s = _fix_acronyms(s)
        return s
    return s

def program_name(m: dict) -> str:
    # prefer real CSV fields before URL slug
    candidates = [
        m.get("name"), m.get("title"), m.get("program"), m.get("call"),
        m.get("call_title"), m.get("funding_title"), m.get("display_name")
    ]
    url = (m.get("url") or "").strip()
    for c in candidates:
        c = fmt(c, "")
        if c:
            return normalize_program_title(c, url=url)
    # fallback: derive from URL
    if url:
        slug = url.rstrip("/").split("/")[-1]
        return normalize_program_title(slug, url=url) or "Unnamed"
    return "Unnamed"

# --- keyword boost/backfill ---
def _normalize_text(x): 
    return str(x or "").lower()

def keyword_candidates(df, query, dom_pref="", top_n=50):
    if df.empty or not query:
        return pd.DataFrame()
    fields = ["name","title","program","call","description","domain","eligibility","location"]
    toks = set(re.findall(r"\w+", query.lower()))
    dom_tok = set(re.findall(r"\w+", dom_pref.lower())) if dom_pref else set()

    def score_row(r):
        hay = " ".join(_normalize_text(r.get(f, "")) for f in fields)
        score = sum(tok in hay for tok in toks)
        if dom_tok and any(d in hay for d in dom_tok):
            score += 2
        return score

    df2 = df.copy()
    df2["kw_score"] = df2.apply(score_row, axis=1)
    df2 = df2[df2["kw_score"] > 0].sort_values("kw_score", ascending=False).head(top_n)
    return df2

def _key_from_item(item):
    return (item.get("url") or "").strip() or _normalize_text(item.get("name") or item.get("title"))

def hybrid_boost(semantic_matches, df_full, query, need_eur=0, dom_pref="", user_loc="", want=8):
    seen = set(_key_from_item(m) for m in semantic_matches)
    kw = keyword_candidates(df_full, query, dom_pref=dom_pref, top_n=50)
    now_ts = pd.Timestamp.now(tz="UTC")

    additions = []
    for _, r in kw.iterrows():
        item = r.to_dict()
        k = _key_from_item(item)
        if k in seen:
            continue
        deadline_val = item.get("deadline")
        item["deadline_date"] = pd.to_datetime(safe_parse_deadline(deadline_val), errors="coerce", utc=True)
        if pd.notnull(item["deadline_date"]):
            days = int((item["deadline_date"] - now_ts).days)
            if days < 0:
                continue
            item["days_left"] = days
        item["relevance_score"] = compute_relevance_score(
            item, query, funding_need=need_eur, target_domain=dom_pref, user_location=user_loc
        )
        additions.append(item)

    merged = semantic_matches + additions
    merged = sorted(merged, key=lambda x: x.get("relevance_score", 0), reverse=True)[:want]
    return merged

# =========================
# 4) Retrieval Pipeline + LLM
# =========================
def query_funding_data(query, location="", top_k=8, need_eur=0, dom_pref=""):
    emb = get_embedding(query)
    res = index.query(vector=emb, top_k=top_k, include_metadata=True, namespace=NAMESPACE)
    matches = [m["metadata"] for m in res.get("matches", [])]
    if not matches: return []
    df = pd.DataFrame(matches)
    if "deadline" in df.columns:
        df["deadline"] = df["deadline"].replace(["", "deadline information not found"], np.nan)
        df["deadline_date"] = pd.to_datetime(df["deadline"].apply(safe_parse_deadline), errors="coerce", utc=True)
        now_ts = pd.Timestamp.now(tz="UTC")
        df["days_left"] = (df["deadline_date"] - now_ts).dt.days
        df = df[(df["days_left"].isna()) | (df["days_left"] >= 0)]
        if df.empty: return []
    df["relevance_score"] = df.apply(
        lambda r: compute_relevance_score(r, query, funding_need=need_eur, target_domain=dom_pref, user_location=location),
        axis=1,
    )
    df = df.sort_values("relevance_score", ascending=False).head(top_k)
    return df.to_dict(orient="records")

def build_llm_selection_payload(query, matches, wanted=3):
    items = []
    for i, m in enumerate(matches, 1):
        items.append({
            "id": i,
            "name": m.get("name",""),
            "title": m.get("title",""),
            "domain": m.get("domain",""),
            "description": (m.get("description","") or "")[:800],
            "eligibility": (m.get("eligibility","") or "")[:400],
            "amount": m.get("amount",""),
            "deadline": m.get("deadline",""),
            "location": m.get("location",""),
            "source": m.get("source",""),
            "url": m.get("url",""),
        })
    return {
      "query": query,
      "programs": items,
      "instruction": (
        f"Pick up to {wanted} unique programs by id. DO NOT select duplicates "
        "(same name/source/url). Prefer programs whose text explicitly mentions the user's domain(s), "
        "region, or a funding amount that matches/exceeds the user's need. "
        "For each pick, give a 1–2 sentence reason that cites exact matches "
        "(e.g., 'mentions robotics', 'Rhineland-Palatinate', 'grant up to €500k'). "
        "Return JSON: {\"picks\":[{\"id\":<int>,\"why\":\"...\"}]}. Use only provided programs."
      )
    }

def llm_select_top(client, query, matches, wanted=3):
    try:
        payload = build_llm_selection_payload(query, matches, wanted=wanted)
        res = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a precise funding advisor. Use only the provided programs. Do not invent anything."},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}
            ],
        )
        data = json.loads(res.choices[0].message.content)
        picks = data.get("picks", [])
        ids, reasons = [], {}
        for p in picks:
            try:
                pid = int(p.get("id"))
                if 1 <= pid <= len(matches):
                    ids.append(pid)
                    if p.get("why"): reasons[pid] = str(p["why"])[:300]
            except Exception:
                continue
        ids = ids[:wanted]
        if not ids:
            ids = list(range(1, min(wanted, len(matches)) + 1))
        return ids, reasons
    except Exception:
        ids = list(range(1, min(wanted, len(matches)) + 1))
        return ids, {}

def llm_enrich_picks(client, chosen_ids, matches):
    items = []
    for cid in chosen_ids:
        m = matches[cid - 1]
        items.append({
            "id": cid,
            "name": m.get("name", ""),
            "title": m.get("title", ""),
            "description": (m.get("description", "") or "")[:800],
            "eligibility": (m.get("eligibility", "") or "")[:500],
            "procedure": (m.get("procedure", "") or "")[:500],
            "source": m.get("source", ""),
            "url": m.get("url", ""),
            "location": m.get("location", ""),
            "domain": m.get("domain", ""),
            "deadline": m.get("deadline", ""),
        })

    user_payload = {
        "instruction": (
            "For each program, write a concise 2-sentence brief that summarizes WHAT the program funds and its goal. "
            "Do NOT include eligibility or domain details in the brief; those are separate fields. "
            "Use ONLY the provided text; do not invent facts. "
            "Then propose up to 3 concrete NEXT STEPS explicitly anchored in 'procedure' and/or 'eligibility' if present "
            "(e.g., 'Prepare SME eligibility proof', 'Submit online form as described'). "
            "If those fields are missing, suggest generic steps (visit page, prepare summary). "
            "Each step ≤ 12 words. Do NOT invent URLs or contacts. "
            "Return JSON: {\"items\":[{\"id\":<int>,\"brief\":\"..\",\"next_steps\":[\"..\",\"..\",\"..\"]}]}. "
            "If description is missing, set brief to 'Not specified'."
        ),
        "programs": items
    }

    try:
        res = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Compress and structure text precisely. Use ONLY provided fields; do not invent facts."},
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}
            ],
            max_tokens=600,
        )
        data = json.loads(res.choices[0].message.content)
        out = {}
        for it in data.get("items", []):
            cid = it.get("id")
            if isinstance(cid, int) and cid in chosen_ids:
                brief = str(it.get("brief", "")).strip()
                steps = [str(s).strip() for s in (it.get("next_steps") or []) if str(s).strip()]
                out[cid] = {"brief": brief[:600], "next_steps": steps[:3]}
        return out
    except Exception:
        return {cid: {"brief": "", "next_steps": []} for cid in chosen_ids}

# =========================
# 5) Follow-up understanding (detail mode)
# =========================
ORDINALS = {"first":1,"1st":1,"one":1,"second":2,"2nd":2,"two":2,"third":3,"3rd":3,"three":3,"fourth":4,"4th":4,"four":4,"fifth":5,"5th":5,"five":5}

def find_followup_target(query_text: str, chosen_ids: list[int], matches: list[dict]):
    """Return (orig_id, program_dict) if the user refers to one of the last picks."""
    if not chosen_ids or not matches:
        return None, None
    q = (query_text or "").lower()

    # 1) ordinal like "second one", "2nd", "the 3rd"
    for word, num in ORDINALS.items():
        if re.search(rf"\b{word}\b", q):
            if 1 <= num <= len(chosen_ids):
                orig_id = chosen_ids[num-1]
                return orig_id, matches[orig_id - 1]

    # 2) by name fuzzy
    def norm(x): return re.sub(r"\s+", " ", str(x or "").lower()).strip()
    for cid in chosen_ids:
        m = matches[cid - 1]
        name = norm(program_name(m))
        tokens = [t for t in re.findall(r"\w+", name) if len(t) > 3]
        if tokens and any(t in q for t in tokens):
            return cid, m

    # 3) generic "tell me more" -> default first
    if re.search(r"(tell me more|details|more info|expand|elaborate)", q):
        cid = chosen_ids[0]
        return cid, matches[cid - 1]

    return None, None

# =========================
# 6) Rendering (chat bubbles only)
# =========================
def render_program_bubbles_html(chosen_pairs, reasons_map, query, enrich_map=None):
    def esc(x): return html.escape(str(x))
    domains = [fmt(m.get("domain")) for _, m in chosen_pairs]
    uniq = {d for d in domains if d != "Not specified"}
    show_domain = len(uniq) > 1  # hide if all same

    for rank, (orig_id, m) in enumerate(chosen_pairs, start=1):
        with st.chat_message("assistant"):
            name     = program_name(m)
            reason   = reasons_map.get(orig_id, "")
            brief    = (enrich_map or {}).get(orig_id, {}).get("brief", "").strip() if enrich_map else ""
            desc     = brief or two_sentences(fmt(m.get("description")))
            domain   = fmt(m.get("domain"), "Not specified")
            elig     = two_sentences(fmt(m.get("eligibility")))
            amount   = fmt(m.get("amount"))
            deadline = deadline_with_badge(m)
            loc_txt  = fmt(m.get("location"))
            contact  = fmt(m.get("contact"))
            src      = fmt(m.get("source"), "N/A")

            html_block = f"""
<div>
  <strong>{esc(rank)}. {esc(name)}</strong><br><br>
  {f"<strong>Why it fits:</strong> {esc(reason)}<br><br>" if reason else ""}
  <strong>Description:</strong> {esc(desc)}<br>
  {f"<strong>Domain:</strong> {esc(domain)}<br>" if show_domain else ""}
  <strong>Eligibility:</strong> {esc(elig)}<br>
  <strong>Amount:</strong> {esc(amount)}<br>
  <strong>Deadline:</strong> {esc(deadline)}<br>
  <strong>Location:</strong> {esc(loc_txt)}<br>
  <strong>Contact:</strong> {esc(contact)}<br>
  <strong>Source:</strong> {esc(src)}<br>
</div>
"""
            st.markdown(html_block, unsafe_allow_html=True)

            # Next Steps (always inject official page link if URL present)
            steps = (enrich_map or {}).get(orig_id, {}).get("next_steps", []) if enrich_map else []
            url = (m.get("url") or "").strip()
            if url:
                url_step = f'Visit the <a href="{html.escape(url)}" target="_blank">official page</a>'
                if not any("official page" in str(s).lower() for s in steps):
                    steps = [url_step] + steps
            if not steps:
                steps = steps_from_dataset(m)

            st.markdown("**Next Steps:**", unsafe_allow_html=True)
            for s in steps[:3]:
                st.markdown(f"- {s}", unsafe_allow_html=True)

            # Draft buttons
            gen_key = f"gen-{orig_id}"
            dl_key  = f"dl-{orig_id}"
            if st.button("📝 Generate Draft", key=gen_key):
                try:
                    qtext = query or st.session_state.get("last_query") or ""
                    buf = build_application_docx(m, qtext)
                    st.session_state.draft_blobs[orig_id] = buf.getvalue()
                    st.success("Draft ready ↓")
                except Exception as e:
                    st.error(f"Couldn't generate the DOCX: {e}")

            if orig_id in st.session_state.draft_blobs:
                st.download_button(
                    label="⬇️ Download .docx",
                    data=st.session_state.draft_blobs[orig_id],
                    file_name=f"application_{rank}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=dl_key
                )

# --- DOCX builder ---
def build_application_docx(program, user_query):
    if not HAS_DOCX:
        raise RuntimeError(
            f"python-docx is not available. Install it with: pip install python-docx "
            f"(import error: {DOCX_IMPORT_ERR})"
        )

    doc = Document()

    def add_heading(text, level=0): doc.add_heading(text, level=level)
    def add_para(text=""): doc.add_paragraph(str(text) if text is not None else "")

    d = program or {}
    add_heading("Grant Application Draft", 0)
    add_para(f"Program: {program_name(d)}")
    add_para(f"Contact: {fmt(d.get('contact'))}")        # Contact above Source
    add_para(f"Source: {fmt(d.get('source'),'N/A')}")
    add_para(f"Amount: {fmt(d.get('amount'))}")
    add_para(f"Deadline: {fmt(d.get('deadline'))}")
    add_para(f"Location: {fmt(d.get('location'))}")
    add_para()
    add_heading("Official Page", 1); add_para(fmt(d.get('url'), "Not specified"))
    add_heading("Company Summary", 1); add_para(user_query or "—")
    add_heading("Why This Fits", 1); add_para(f"Domain match: {fmt(d.get('domain'),'N/A')}"); add_para(f"Eligibility: {fmt(d.get('eligibility'))}")
    add_heading("Next Steps", 1); add_para(fmt(d.get("procedure"), "Not specified"))
    for it in [
        "Collect mandatory docs (CV, financials, project plan)",
        "Draft project timeline & budget",
        "Contact the program office for clarifications",
    ]:
        add_para(f"• {it}")
    add_heading("Checklist", 1)
    for it in ["Eligibility confirmed","Budget aligned","Deadlines planned","All docs prepared"]:
        add_para(f"☐ {it}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# =========================
# 7) Main Chat Flow (with follow-up handling)
# =========================
# Show prior transcript
for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Input
query = st.chat_input("Describe your company or ask a follow-up…") or st.session_state.get("generated_query", "")

def build_chosen_pairs(chosen_ids, matches):
    return [(cid, matches[cid - 1]) for cid in (chosen_ids or []) if 1 <= cid <= len(matches)]

def run_full_pipeline(user_query: str):
    base = query_funding_data(user_query, location="", top_k=8, need_eur=0, dom_pref="")
    matches = hybrid_boost(base, df_full, user_query, need_eur=0, dom_pref="", user_loc="", want=8)
    wanted = 3
    chosen_ids, reasons = llm_select_top(client, user_query, matches, wanted=wanted)
    enrich_map = llm_enrich_picks(client, chosen_ids, matches)
    # Persist
    st.session_state.last_query = user_query
    st.session_state.last_matches = matches
    st.session_state.last_chosen_ids = chosen_ids
    st.session_state.last_reasons = reasons
    st.session_state.last_enrich = enrich_map
    return build_chosen_pairs(chosen_ids, matches), reasons, enrich_map

def render_detail_for_followup(target_id: int, program: dict, user_query: str):
    """Render a single, expanded bubble for the selected program (no re-listing)."""
    chosen_pairs = [(target_id, program)]
    reasons = st.session_state.last_reasons or {}
    enrich_map = st.session_state.last_enrich or {}
    render_program_bubbles_html(chosen_pairs, reasons, user_query, enrich_map=enrich_map)

if query:
    with st.chat_message("user"):
        st.markdown(query)
    st.session_state.chat_history.append({"role": "user", "content": query})

    # If we have previous picks, try follow-up first
    if st.session_state.last_matches and st.session_state.last_chosen_ids:
        orig_id, prog = find_followup_target(query, st.session_state.last_chosen_ids, st.session_state.last_matches)
        if orig_id and prog:
            with st.spinner("Pulling details…"):
                render_detail_for_followup(orig_id, prog, st.session_state.last_query or query)
        else:
            with st.spinner("Finding and ranking programs…"):
                pairs, reasons, enrich = run_full_pipeline(query)
                render_program_bubbles_html(pairs, reasons, query, enrich_map=enrich)
    else:
        with st.spinner("Finding and ranking programs…"):
            pairs, reasons, enrich = run_full_pipeline(query)
            render_program_bubbles_html(pairs, reasons, query, enrich_map=enrich)
