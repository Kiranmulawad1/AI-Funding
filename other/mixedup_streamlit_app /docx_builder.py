# docx_builder.py
import io
try:
    from docx import Document
    HAS_DOCX = True
    DOCX_IMPORT_ERR = None
except Exception as _e:
    HAS_DOCX = False
    DOCX_IMPORT_ERR = _e

from utils import fmt, program_name

def build_application_docx(program: dict, user_query: str) -> io.BytesIO:
    if not HAS_DOCX:
        raise RuntimeError(
            f"python-docx is not available. Install it with: pip install python-docx "
            f"(import error: {DOCX_IMPORT_ERR})"
        )

    doc = Document()

    def add_heading(text, level=0): doc.add_heading(text, level=level)
    def add_para(text=""): doc.add_paragraph(str(text) if text is not None else "")

    d = program or {}
    add_heading("Grant Application Draft", 0)
    add_para(f"Program: {program_name(d)}")
    add_para(f"Contact: {fmt(d.get('contact'))}")
    add_para(f"Source: {fmt(d.get('source'),'N/A')}")
    add_para(f"Amount: {fmt(d.get('amount'))}")
    add_para(f"Deadline: {fmt(d.get('deadline'))}")
    add_para(f"Location: {fmt(d.get('location'))}")
    add_para()
    add_heading("Official Page", 1); add_para(fmt(d.get('url'), "Not specified"))
    add_heading("Company Summary", 1); add_para(user_query or "—")
    add_heading("Why This Fits", 1); add_para(f"Domain match: {fmt(d.get('domain'),'N/A')}"); add_para(f"Eligibility: {fmt(d.get('eligibility'))}")
    add_heading("Next Steps", 1); add_para(fmt(d.get("procedure"), "Not specified"))
    for it in [
        "Collect mandatory docs (CV, financials, project plan)",
        "Draft project timeline & budget",
        "Contact the program office for clarifications",
    ]:
        add_para(f"• {it}")
    add_heading("Checklist", 1)
    for it in ["Eligibility confirmed","Budget aligned","Deadlines planned","All docs prepared"]:
        add_para(f"☐ {it}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
