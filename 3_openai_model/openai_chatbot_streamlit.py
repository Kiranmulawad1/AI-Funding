import os
import re
import io
import time
import json
import pandas as pd
import numpy as np
import streamlit as st
from dateutil import parser
from dotenv import load_dotenv
from openai import OpenAI
from pinecone import Pinecone
import fitz  # PyMuPDF

# =========================
# 0) Setup & Environment
# =========================
load_dotenv()

OPENAI_API_KEY   = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENV     = os.getenv("PINECONE_ENV")
INDEX_NAME       = os.getenv("PINECONE_INDEX_NAME", "funding-search")
NAMESPACE        = os.getenv("PINECONE_NAMESPACE", "openai-v3")
FUNDING_CSV_PATH = os.getenv("FUNDING_CSV_PATH", "./data/merged_funding_data.csv")

required = {
    "OPENAI_API_KEY": OPENAI_API_KEY,
    "PINECONE_API_KEY": PINECONE_API_KEY,
    "PINECONE_ENV": PINECONE_ENV,
}
st.set_page_config(page_title="Smart Funding Chatbot", layout="centered")
if missing := [k for k, v in required.items() if not v]:
    st.error(f"Missing environment variables: {', '.join(missing)}")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)
pc = Pinecone(api_key=PINECONE_API_KEY, environment=PINECONE_ENV)

@st.cache_resource(show_spinner=False)
def get_index(_pc, name: str):
    return _pc.Index(name)

@st.cache_data(show_spinner=False)
def load_full_df(path: str):
    if os.path.exists(path):
        return pd.read_csv(path)
    return pd.DataFrame()

index = get_index(pc, INDEX_NAME)
df_full = load_full_df(FUNDING_CSV_PATH)

# =========================
# 1) Session State & Defaults
# =========================
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "last_matches" not in st.session_state:
    st.session_state.last_matches = []
if "generated_query" not in st.session_state:
    st.session_state.generated_query = ""
if "draft_blobs" not in st.session_state:
    st.session_state.draft_blobs = {}  # {orig_id: bytes}
if "last_llm_picks" not in st.session_state:
    st.session_state.last_llm_picks = []  # [(orig_id, reason)]

RECOMMENDATIONS_TO_SHOW = 3  # we’ll still retrieve ~3; LLM trims to top 2–3
DEFAULT_NEED_EUR = 0
DEFAULT_LOC = ""
DEFAULT_DOMAIN_PREF = ""

# =========================
# 2) Utilities
# =========================
def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text.strip()

def summarize_text_with_gpt(raw_text: str) -> str:
    prompt = f"""
Summarize this company profile into 2–3 lines for public funding discovery in Germany.
Focus on domain, goals, and funding need.

---
{raw_text[:6000]}
---
"""
    res = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return res.choices[0].message.content.strip()

def get_embedding(text: str):
    return client.embeddings.create(input=[text], model="text-embedding-3-small").data[0].embedding

def safe_parse_deadline(deadline_str):
    try:
        if pd.isna(deadline_str) or str(deadline_str).strip() == "":
            return None
        return parser.parse(str(deadline_str), dayfirst=True, fuzzy=True)
    except Exception:
        return None

def compute_relevance_score(row, query, funding_need=0, target_domain="", user_location=""):
    get = (row.get if isinstance(row, dict) else lambda k, d=None: row.get(k, d))
    score = 0.0
    if target_domain and target_domain.lower() in str(get("domain","")).lower():
        score += 0.4
    try:
        nums = re.findall(r"\d+", str(get("amount","0")))
        amount_val = int(nums[-1]) if nums else 0
        if funding_need and amount_val >= funding_need:
            score += 0.3
    except Exception:
        pass
    dd = get("deadline_date")
    if isinstance(dd, pd.Timestamp) and pd.notnull(dd) and dd >= pd.Timestamp.now(tz="UTC"):
        score += 0.2
    tokens = re.findall(r"\w+", query.lower())
    if any(tok in str(get("description","")).lower() for tok in tokens):
        score += 0.1
    if user_location and user_location.lower() in str(get("location","")).lower():
        score += 0.1
    return round(score * 100)

def days_left_utc(deadline_str):
    d = safe_parse_deadline(deadline_str)
    if d is None:
        return None
    d_utc = pd.to_datetime(d, errors="coerce", utc=True)
    if pd.isna(d_utc):
        return None
    return (d_utc - pd.Timestamp.now(tz="UTC")).days

# ---- Hybrid keyword boost/backfill ----
def _normalize_text(x): 
    return str(x or "").lower()

def keyword_candidates(df, query, dom_pref="", top_n=20):
    if df.empty or not query:
        return pd.DataFrame()
    fields = ["name","description","domain","eligibility","location"]
    toks = set(re.findall(r"\w+", query.lower()))
    dom_tok = set(re.findall(r"\w+", dom_pref.lower())) if dom_pref else set()

    def score_row(r):
        hay = " ".join(_normalize_text(r.get(f, "")) for f in fields)
        score = sum(tok in hay for tok in toks)
        if dom_tok and any(d in hay for d in dom_tok):
            score += 2
        return score

    df2 = df.copy()
    df2["kw_score"] = df2.apply(score_row, axis=1)
    df2 = df2[df2["kw_score"] > 0].sort_values("kw_score", ascending=False).head(top_n)
    return df2

def _key_from_item(item):
    return (item.get("url") or "").strip() or _normalize_text(item.get("name"))

def hybrid_boost(semantic_matches, df_full, query, need_eur=0, dom_pref="", user_loc="", want=3):
    seen = set(_key_from_item(m) for m in semantic_matches)
    kw = keyword_candidates(df_full, query, dom_pref=dom_pref, top_n=50)

    additions = []
    for _, r in kw.iterrows():
        item = r.to_dict()
        k = _key_from_item(item)
        if k in seen:
            continue
        deadline_val = item.get("deadline")
        item["deadline_date"] = pd.to_datetime(safe_parse_deadline(deadline_val), errors="coerce", utc=True)
        item["relevance_score"] = compute_relevance_score(
            item, query, funding_need=need_eur, target_domain=dom_pref, user_location=user_loc
        )
        additions.append(item)

    merged = semantic_matches + additions
    merged = sorted(merged, key=lambda x: x.get("relevance_score", 0), reverse=True)[:want]
    return merged

# ---- DOCX application draft ----
def build_application_docx(program, user_query):
    from docx import Document  # pip install python-docx
    doc = Document()

    def add_heading(text, level=0):
        doc.add_heading(text, level=level)
    def add_para(text=""):
        doc.add_paragraph(str(text) if text is not None else "")

    d = program or {}
    add_heading("Grant Application Draft", level=0)
    add_para(f"Program: {d.get('name','Unnamed')}")
    add_para(f"Source: {d.get('source','N/A')}")
    add_para(f"URL: {d.get('url','')}")
    add_para(f"Amount: {d.get('amount','Not specified')}")
    add_para(f"Deadline: {d.get('deadline','Not specified')}")
    add_para(f"Location: {d.get('location','Not specified')}")
    add_para()

    add_heading("Company Summary", 1); add_para(user_query or "—")
    add_heading("Why This Fits", 1)
    add_para(f"Domain match: {d.get('domain','N/A')}")
    add_para(f"Eligibility: {d.get('eligibility','Not specified')}")
    add_heading("Next Steps", 1)
    add_para(d.get("procedure","Not specified"))
    for it in ["Collect mandatory docs (CV, financials, project plan)",
               "Draft project timeline & budget",
               "Contact the program office for clarifications"]:
        add_para(f"• {it}")
    add_heading("Contacts", 1); add_para(d.get("contact","Not specified"))
    add_heading("Checklist", 1)
    for it in ["Eligibility confirmed","Budget aligned","Deadlines planned","All docs prepared"]:
        add_para(f"☐ {it}")

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio

# ---- Retrieval (semantic) ----
def query_funding_data(query, location="", top_k=3, need_eur=0, dom_pref=""):
    emb = get_embedding(query)
    res = index.query(vector=emb, top_k=top_k, include_metadata=True, namespace=NAMESPACE)
    matches = [m["metadata"] for m in res.get("matches", [])]
    if not matches:
        return []

    df = pd.DataFrame(matches)
    if "deadline" in df.columns:
        df["deadline"] = df["deadline"].replace(["", "deadline information not found"], np.nan)
        df["deadline_date"] = pd.to_datetime(df["deadline"].apply(safe_parse_deadline), errors="coerce", utc=True)
        now_ts = pd.Timestamp.now(tz="UTC")
        df["days_left"] = (df["deadline_date"] - now_ts).dt.days
        df = df[(df["days_left"].isna()) | (df["days_left"] >= 0)]
        if df.empty:
            return []

    df["relevance_score"] = df.apply(
        lambda r: compute_relevance_score(
            r, query, funding_need=need_eur, target_domain=dom_pref, user_location=location
        ),
        axis=1,
    )
    df = df.sort_values("relevance_score", ascending=False).head(top_k)
    return df.to_dict(orient="records")

# ---- Build compact JSON prompt for LLM selection ----
def build_llm_selection_payload(query, matches):
    items = []
    for i, m in enumerate(matches, 1):
        items.append({
            "id": i,
            "name": m.get("name",""),
            "domain": m.get("domain",""),
            "description": (m.get("description","") or "")[:500],
            "eligibility": (m.get("eligibility","") or "")[:300],
            "amount": m.get("amount",""),
            "deadline": m.get("deadline",""),
            "location": m.get("location",""),
        })
    return {
      "query": query,
      "programs": items,
      "instruction": (
        "Pick the best 2–3 programs by id. Optimize for domain fit, maturity stage, and amount/eligibility. "
        "If info is missing, still choose the closest fit. Return JSON: "
        '{"picks":[{"id":<int>,"why":"1-2 sentences"}]}'
      )
    }

def llm_select_top(client, query, matches):
    try:
        payload = build_llm_selection_payload(query, matches)
        res = client.chat.completions.create(
            model="gpt-4-turbo",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a precise funding advisor. Use only the provided programs. Do not invent anything."},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}
            ],
        )
        data = json.loads(res.choices[0].message.content)
        picks = data.get("picks", [])
        ids = []
        reasons = {}
        for p in picks:
            try:
                pid = int(p.get("id"))
                if 1 <= pid <= len(matches):
                    ids.append(pid)
                    if p.get("why"): reasons[pid] = str(p["why"])[:300]
            except Exception:
                continue
        # ensure 2–3 items max
        ids = ids[:3]
        if not ids:
            ids = list(range(1, min(3, len(matches)) + 1))
        return ids, reasons
    except Exception:
        # fallback to top-N if any error
        ids = list(range(1, min(3, len(matches)) + 1))
        return ids, {}

# ---- Render one bubble per program (with its own button + reason) ----
def render_program_bubbles(matches, chosen_ids, reasons_map, query, delay=0.18):
    chosen = [(cid, matches[cid-1]) for cid in chosen_ids]
    placeholders = [st.empty() for _ in range(len(chosen))]
    for rank, (ph, item) in enumerate(zip(placeholders, chosen), start=1):
        orig_id, m = item
        with ph.container():
            with st.chat_message("assistant"):
                name = m.get("name", "Unnamed")
                amount = m.get("amount", "Not specified")
                deadline = m.get("deadline", "Not specified")
                loc_txt = m.get("location", "Not specified")
                src = m.get("source", "N/A")
                url = m.get("url", "")
                reason = reasons_map.get(orig_id, "")

                md = f"**{rank}. {name}**"
                if reason:
                    md += f"\n\n**Why it fits:** {reason}"
                md += (f"\n\n**Amount:** {amount} | **Deadline:** {deadline} | **Location:** {loc_txt}"
                       f"\n**Source:** {src}" + (f"\n[Open Call]({url})" if url else ""))
                st.markdown(md)

                gen_key = f"gen-{orig_id}"
                dl_key = f"dl-{orig_id}"
                if st.button("📝 Generate Draft", key=gen_key):
                    buf = build_application_docx(m, query)
                    st.session_state.draft_blobs[orig_id] = buf.getvalue()
                    st.success("Draft ready ↓")
                if orig_id in st.session_state.draft_blobs:
                    st.download_button(
                        "⬇️ Download DOCX",
                        st.session_state.draft_blobs[orig_id],
                        file_name=f"application_{rank}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=dl_key
                    )
        time.sleep(delay)

# =========================
# 3) Sidebar (PDF only)
# =========================
st.sidebar.subheader("📄 Upload Company Profile (Optional)")
uploaded_pdf = st.sidebar.file_uploader("Upload PDF", type=["pdf"])
if uploaded_pdf:
    raw_text = extract_text_from_pdf(uploaded_pdf)
    summary_query = summarize_text_with_gpt(raw_text)  # gpt-3.5 for summarization
    st.session_state.generated_query = summary_query  # hidden
    st.sidebar.success("✅ Profile processed. Using it to search.")

# =========================
# 4) Chat UI (LLM-curated shortlist, progressive bubbles)
# =========================
for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

query = st.chat_input("Describe your company or upload a profile PDF…") or st.session_state.get("generated_query", "")

if query:
    with st.chat_message("user"):
        st.markdown(query)
    st.session_state.chat_history.append({"role": "user", "content": query})

    # Retrieval under spinner
    with st.spinner("Finding and ranking programs…"):
        is_followup = len(st.session_state.last_matches) > 0 and len(st.session_state.chat_history) > 1
        if is_followup:
            matches = st.session_state.last_matches
        else:
            base = query_funding_data(
                query,
                location=DEFAULT_LOC,
                top_k=RECOMMENDATIONS_TO_SHOW,
                need_eur=DEFAULT_NEED_EUR,
                dom_pref=DEFAULT_DOMAIN_PREF
            )
            matches = hybrid_boost(
                base, df_full, query,
                need_eur=DEFAULT_NEED_EUR,
                dom_pref=DEFAULT_DOMAIN_PREF,
                user_loc=DEFAULT_LOC,
                want=RECOMMENDATIONS_TO_SHOW
            )
            st.session_state.last_matches = matches

        # LLM picks top 2–3 by ID (guardrailed JSON)
        chosen_ids, reasons = llm_select_top(client, query, matches)
        st.session_state.last_llm_picks = [(cid, reasons.get(cid, "")) for cid in chosen_ids]

    if not matches:
        with st.chat_message("assistant"):
            st.info("No suitable programs found. Try a broader description or upload a profile PDF.")
        st.session_state.chat_history.append({"role": "assistant", "content": "No suitable programs found."})
    else:
        # Progressive assistant bubbles for the chosen 2–3
        render_program_bubbles(matches, chosen_ids, reasons, query, delay=0.2)
