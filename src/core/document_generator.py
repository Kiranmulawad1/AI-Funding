import io
from docx import Document

def build_draft_prompt(profile, metadata):
    return f"""You are a professional grant writer AI assistant.
Use the following company profile and the funding program details to generate a complete funding application draft.

## Company Profile
- Company Name: {profile.get("company_name", "Not specified")}
- Location: {profile.get("location", "Not specified")}
- Industry: {profile.get("industry", "Not specified")}
- Goals: {profile.get("goals", "Not specified")}
- Project Idea: {profile.get("project_idea", "Not specified")}
- Funding Need: {profile.get("funding_need", "Not specified")}

## Funding Program
- Name: {metadata.get("name", "Not specified")}
- Amount: {metadata.get("amount", "Not specified")}
- Deadline: {metadata.get("deadline", "Not specified")}
- Eligibility: {metadata.get("eligibility", "Not specified")}

## Format Output as:
1. Executive Summary
2. Project Description
3. Technical Approach
4. Budget Overview
5. Expected Outcomes
6. Relevance to Program

Be professional and compelling."""

def generate_funding_draft(metadata, profile, llm_client, content: str = None):
    if content:
        # If content provides, use it directly (bypass LLM)
        draft_text = content
    else:
        # Otherwise generate it
        prompt = build_draft_prompt(profile, metadata)
        
        response = llm_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        
        draft_text = response.choices[0].message.content.strip()
    
    doc = Document()
    doc.add_heading("Funding Application Draft", 0)
    
    for para in draft_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio