# funding_draft_generator.py

import io
from docx import Document
from utils import present, program_name

def build_draft_prompt(program: dict, query: str) -> str:
    return f"""
You are an assistant helping with funding applications in Germany.

Based on the following company info and grant metadata, draft a short application template
that can be edited and submitted to the funder. Focus on clarity, structure, and next steps.

### Company Info:
{query}

### Grant Info:
Program: {program_name(program)}
Domain: {present(program.get("domain"))}
Eligibility: {present(program.get("eligibility"))}
Funding Amount: {present(program.get("amount"))}
Deadline: {present(program.get("deadline"))}
Location: {present(program.get("location"))}
Contact: {present(program.get("contact"))}
Procedure: {present(program.get("procedure"))}

Give the result in structured paragraphs with headings: Introduction, Why This Grant, Next Steps, Checklist.
"""

def generate_funding_draft(program: dict, query: str, llm_client) -> io.BytesIO:
    prompt = build_draft_prompt(program, query)
    
    response = llm_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )

    draft_text = response.choices[0].message.content.strip()

    doc = Document()
    doc.add_heading("Funding Application Draft", 0)

    for section in draft_text.split("\n\n"):
        if section.strip():
            doc.add_paragraph(section.strip())

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
